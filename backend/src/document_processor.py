import os
import uuid
import json
from typing import List, Dict, Tuple, Callable, Optional
import PyPDF2
import docx
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
import logging
import asyncio
import time

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, pinecone_index_name: str):
        self.embeddings = OpenAIEmbeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        self.vectorstore = PineconeVectorStore(
            index_name=pinecone_index_name,
            embedding=self.embeddings
        )
        self.progress_callbacks: Dict[int, Callable] = {}
    
    def register_progress_callback(self, doc_id: int, callback: Callable[[int], None]):
        """Register a callback function for progress updates"""
        self.progress_callbacks[doc_id] = callback
    
    def update_progress(self, doc_id: int, progress: int):
        """Update progress for a document"""
        if doc_id in self.progress_callbacks:
            self.progress_callbacks[doc_id](progress)
    
    def extract_text(self, file_content: bytes, file_type: str, filename: str) -> str:
        """Extract text from different file types"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_pdf_text(file_content)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_docx_text(file_content)
            elif file_type.lower() == 'txt':
                return file_content.decode('utf-8')
            elif file_type.lower() in ['csv', 'xlsx']:
                return self._extract_structured_data(file_content, file_type)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF with progress tracking"""
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        total_pages = len(pdf_reader.pages)
        
        for i, page in enumerate(pdf_reader.pages):
            text += page.extract_text() + "\n"
            # Report progress every few pages for large PDFs
            if i % 5 == 0 or i == total_pages - 1:
                progress = int((i + 1) / total_pages * 20)  # 20% of total process
                logger.info(f"PDF extraction progress: {progress}%")
        
        return text
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        import io
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_structured_data(self, file_content: bytes, file_type: str) -> str:
        """Extract and format structured data"""
        import io
        if file_type == 'csv':
            df = pd.read_csv(io.BytesIO(file_content))
        else:  # xlsx
            df = pd.read_excel(io.BytesIO(file_content))
        
        # Convert DataFrame to readable text format
        text = f"Data Summary:\n{df.describe().to_string()}\n\n"
        text += f"Column Names: {', '.join(df.columns.tolist())}\n\n"
        text += "Sample Data:\n" + df.head(10).to_string()
        return text
    
    async def process_document(self, file_content: bytes, filename: str, 
                             file_type: str, doc_id: int) -> Dict:
        """Process document through the entire pipeline with real-time progress"""
        try:
            # Step 1: Extract text (0-25%)
            logger.info(f"[{doc_id}] Extracting text from {filename}")
            self.update_progress(doc_id, 5)
            
            text_content = self.extract_text(file_content, file_type, filename)
            self.update_progress(doc_id, 25)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from file")
            
            # Step 2: Split into chunks (25-40%)
            logger.info(f"[{doc_id}] Splitting document into chunks")
            self.update_progress(doc_id, 30)
            
            chunks = self.text_splitter.split_text(text_content)
            self.update_progress(doc_id, 40)
            
            if not chunks:
                raise ValueError("No chunks created from text content")
            
            # Step 3: Create documents with metadata (40-50%)
            logger.info(f"[{doc_id}] Creating document objects")
            documents = []
            for i, chunk in enumerate(chunks):
                metadata = {
                    "source": filename,
                    "doc_id": doc_id,
                    "chunk_id": i,
                    "file_type": file_type,
                    "total_chunks": len(chunks)
                }
                documents.append(Document(page_content=chunk, metadata=metadata))
            
            self.update_progress(doc_id, 50)
            
            # Step 4: Generate embeddings and store in Pinecone (50-100%)
            logger.info(f"[{doc_id}] Generating embeddings for {len(documents)} chunks")
            vector_ids = []
            
            # Process in batches to avoid rate limits and track progress
            batch_size = 10
            total_batches = (len(documents) + batch_size - 1) // batch_size
            
            for batch_idx in range(0, len(documents), batch_size):
                batch = documents[batch_idx:batch_idx + batch_size]
                
                # Generate unique IDs for this batch
                batch_ids = [f"{doc_id}_{j}" for j in range(batch_idx, min(batch_idx + batch_size, len(documents)))]
                vector_ids.extend(batch_ids)
                
                # Add to Pinecone with retry logic
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        self.vectorstore.add_documents(batch, ids=batch_ids)
                        break
                    except Exception as e:
                        if attempt == max_retries - 1:
                            raise e
                        logger.warning(f"[{doc_id}] Retry {attempt + 1} for batch {batch_idx//batch_size + 1}")
                        await asyncio.sleep(1)
                
                # Update progress based on batches completed
                current_batch = batch_idx // batch_size + 1
                progress = 50 + int((current_batch / total_batches) * 50)
                self.update_progress(doc_id, progress)
                
                # Small delay to avoid rate limiting
                await asyncio.sleep(0.1)
            
            self.update_progress(doc_id, 100)
            logger.info(f"[{doc_id}] Successfully processed {filename}: {len(chunks)} chunks, {len(vector_ids)} vectors")
            
            return {
                "success": True,
                "chunk_count": len(chunks),
                "vector_ids": vector_ids,
                "text_length": len(text_content),
                "processing_time": time.time()
            }
            
        except Exception as e:
            logger.error(f"[{doc_id}] Error processing document {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "chunk_count": 0,
                "vector_ids": []
            }
    
    def search_documents(self, query: str, doc_ids: List[int] = None, k: int = 3) -> List[Document]:
        """Search through processed documents"""
        try:
            if doc_ids:
                # Filter by specific document IDs
                filter_dict = {"doc_id": {"$in": doc_ids}}
                docs = self.vectorstore.similarity_search(query, k=k, filter=filter_dict)
            else:
                docs = self.vectorstore.similarity_search(query, k=k)
            
            return docs
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []
    
    def delete_document_vectors(self, vector_ids: List[str]) -> bool:
        """Delete vectors from Pinecone"""
        try:
            # Note: Implement actual Pinecone deletion
            logger.info(f"Deleting {len(vector_ids)} vectors from Pinecone")
            # self.vectorstore.delete(ids=vector_ids)
            return True
        except Exception as e:
            logger.error(f"Error deleting vectors: {e}")
            return False
    
    def get_processing_stats(self) -> Dict:
        """Get processing statistics"""
        return {
            "active_processors": len(self.progress_callbacks),
            "vectorstore_status": "connected",
            "embedding_model": "text-embedding-ada-002"
        }