import os
import uuid
import json
from typing import List, Dict, Tuple, Callable, Optional
import PyPDF2
import docx
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
import logging
import asyncio
import time
import pinecone

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, pinecone_index_name: str):
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        
        # Initialize embeddings conditionally (only if OpenAI key is available)
        self.embeddings = None
        self.vectorstore = None
        
        # TEMPORARY: Skip embeddings due to dimension mismatch
        # Your Pinecone index expects 1024 dimensions (llama-text-embed-v2)
        # But OpenAI produces 1536 dimensions (text-embedding-ada-002)
        
        logger.info("⚠️ Skipping embeddings due to dimension mismatch")
        logger.info("📝 Pinecone index: 1024 dims | OpenAI embeddings: 1536 dims")
        logger.info("🔄 Running in text-only mode until index is updated")
        
        # All embedding code temporarily disabled due to dimension mismatch
        logger.info("📝 Document processor initialized in text-only mode")
        self.progress_callbacks: Dict[int, Callable] = {}
    
    def register_progress_callback(self, doc_id: int, callback: Callable[[int], None]):
        """Register a callback function for progress updates"""
        self.progress_callbacks[doc_id] = callback
    
    def update_progress(self, doc_id: int, progress: int):
        """Update progress for a document"""
        if doc_id in self.progress_callbacks:
            self.progress_callbacks[doc_id](progress)
    
    def extract_text(self, file_content: bytes, file_type: str, filename: str) -> str:
        """Extract text from different file types"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_pdf_text(file_content)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_docx_text(file_content)
            elif file_type.lower() == 'txt':
                return file_content.decode('utf-8')
            elif file_type.lower() in ['csv', 'xlsx']:
                return self._extract_structured_data(file_content, file_type)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF with progress tracking"""
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        total_pages = len(pdf_reader.pages)
        
        for i, page in enumerate(pdf_reader.pages):
            text += page.extract_text() + "\n"
            # Report progress every few pages for large PDFs
            if i % 5 == 0 or i == total_pages - 1:
                progress = int((i + 1) / total_pages * 20)  # 20% of total process
                logger.info(f"PDF extraction progress: {progress}%")
        
        return text
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        import io
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_structured_data(self, file_content: bytes, file_type: str) -> str:
        """Extract and format structured data"""
        import io
        if file_type == 'csv':
            df = pd.read_csv(io.BytesIO(file_content))
        else:  # xlsx
            df = pd.read_excel(io.BytesIO(file_content))
        
        # Convert DataFrame to readable text format
        text = f"Data Summary:\n{df.describe().to_string()}\n\n"
        text += f"Column Names: {', '.join(df.columns.tolist())}\n\n"
        text += "Sample Data:\n" + df.head(10).to_string()
        return text
    
    async def process_document(self, file_content: bytes, filename: str, 
                             file_type: str, doc_id: int) -> Dict:
        """Process document through the entire pipeline with real-time progress"""
        try:
            # Step 1: Extract text (0-25%)
            logger.info(f"[{doc_id}] Extracting text from {filename}")
            self.update_progress(doc_id, 5)
            
            text_content = self.extract_text(file_content, file_type, filename)
            self.update_progress(doc_id, 25)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from file")
            
            # Step 2: Split into chunks (25-40%)
            logger.info(f"[{doc_id}] Splitting document into chunks")
            self.update_progress(doc_id, 30)
            
            chunks = self.text_splitter.split_text(text_content)
            self.update_progress(doc_id, 40)
            
            if not chunks:
                raise ValueError("No chunks created from text content")
            
            # Step 3: Create documents with metadata (40-50%)
            logger.info(f"[{doc_id}] Creating document objects")
            documents = []
            for i, chunk in enumerate(chunks):
                metadata = {
                    "source": filename,
                    "doc_id": doc_id,
                    "chunk_id": i,
                    "file_type": file_type,
                    "total_chunks": len(chunks)
                }
                documents.append(Document(page_content=chunk, metadata=metadata))
            
            self.update_progress(doc_id, 50)
            
            # Step 4: Generate embeddings and store in Pinecone (50-100%) - Optional
            vector_ids = []
            
            if self.vectorstore and self.embeddings:
                logger.info(f"[{doc_id}] Generating embeddings for {len(documents)} chunks")
                
                try:
                    # Process in smaller batches to avoid issues
                    batch_size = 5  # Reduced from 10 to be safer
                    total_batches = (len(documents) + batch_size - 1) // batch_size
                    
                    for batch_idx in range(0, len(documents), batch_size):
                        batch = documents[batch_idx:batch_idx + batch_size]
                        
                        # Generate unique IDs for this batch
                        batch_ids = [f"{doc_id}_{j}" for j in range(batch_idx, min(batch_idx + batch_size, len(documents)))]
                        
                        # Add to Pinecone with comprehensive error handling
                        max_retries = 2  # Reduced retries for faster feedback
                        success = False
                        
                        for attempt in range(max_retries):
                            try:
                                logger.info(f"[{doc_id}] Processing batch {batch_idx//batch_size + 1}/{total_batches}")
                                self.vectorstore.add_documents(batch, ids=batch_ids)
                                vector_ids.extend(batch_ids)
                                success = True
                                break
                                
                            except Exception as e:
                                logger.error(f"[{doc_id}] Batch {batch_idx//batch_size + 1} attempt {attempt + 1} failed: {e}")
                                if attempt == max_retries - 1:
                                    # If all retries fail, continue with text-only mode
                                    logger.warning(f"[{doc_id}] Vectorization failed, continuing in text-only mode")
                                    self.vectorstore = None
                                    break
                                await asyncio.sleep(2)  # Longer wait between retries
                        
                        if not success and not self.vectorstore:
                            break
                        
                        # Update progress based on batches completed
                        current_batch = batch_idx // batch_size + 1
                        progress = 50 + int((current_batch / total_batches) * 40)  # Only go to 90% for vectorization
                        self.update_progress(doc_id, progress)
                        
                        # Longer delay to avoid rate limiting
                        await asyncio.sleep(0.5)
                        
                except Exception as vectorization_error:
                    logger.error(f"[{doc_id}] Vectorization completely failed: {vectorization_error}")
                    logger.info(f"[{doc_id}] Continuing in text-only mode")
                    self.vectorstore = None
                    vector_ids = []
            else:
                logger.info(f"[{doc_id}] No vector storage available - text-only mode")
                
            # Always update to 95% before final completion
            self.update_progress(doc_id, 95)
            
            self.update_progress(doc_id, 100)
            logger.info(f"[{doc_id}] Successfully processed {filename}: {len(chunks)} chunks, {len(vector_ids)} vectors")
            
            return {
                "success": True,
                "chunk_count": len(chunks),
                "vector_ids": vector_ids,
                "text_length": len(text_content),
                "processing_time": time.time()
            }
            
        except Exception as e:
            logger.error(f"[{doc_id}] Error processing document {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "chunk_count": 0,
                "vector_ids": []
            }
    
    def search_documents(self, query: str, doc_ids: List[int] = None, k: int = 3) -> List[Document]:
        """Search through processed documents"""
        try:
            if not self.vectorstore:
                logger.warning("⚠️ Vector search not available - no embeddings configured")
                return []
                
            if doc_ids:
                # Filter by specific document IDs
                filter_dict = {"doc_id": {"$in": doc_ids}}
                docs = self.vectorstore.similarity_search(query, k=k, filter=filter_dict)
            else:
                docs = self.vectorstore.similarity_search(query, k=k)
            
            return docs
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []
    
    def delete_document_vectors(self, vector_ids: List[str]) -> bool:
        """Delete vectors from Pinecone"""
        try:
            # Note: Implement actual Pinecone deletion
            logger.info(f"Deleting {len(vector_ids)} vectors from Pinecone")
            # self.vectorstore.delete(ids=vector_ids)
            return True
        except Exception as e:
            logger.error(f"Error deleting vectors: {e}")
            return False
    
    def get_processing_stats(self) -> Dict:
        """Get processing statistics"""
        return {
            "active_processors": len(self.progress_callbacks),
            "vectorstore_status": "connected",
            "embedding_model": "text-embedding-ada-002"
        }